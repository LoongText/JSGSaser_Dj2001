import jieba
import jieba.analyse
import jieba.posseg
import math
from jsg.settings import PUN_LIST
from tables.models import Projects
from jsg.settings import CLUE_KEYWORDS_LIST
from jsg.settings import HOPE_KEYWORDS_LIST


class TextSummary(object):
    """
    # 提取文本内容的关键信息；
    """
    text = ""
    title = ""
    keywords = list()
    sentences = list()
    summary = list()

    def __init__(self, title, text):
        self.title = title
        self.text = text

    def __split_sentence(self):
        # 通过换行符对文档进行分段
        sections = self.text.split("\n")
        for section in sections:
            if section == "":
                sections.remove(section)

        # 通过分割符对每个段落进行分句
        for i in range(len(sections)):
            section = sections[i]
            text = ""
            k = 0
            for j in range(len(section)):
                char = section[j]
                text = text + char
                if char in ["!",  "。", "？", "；", "?", ";"] or j == len(section)-1:
                    text = text.strip()
                    sentence = dict()
                    sentence["text"] = text
                    sentence["pos"] = dict()
                    sentence["pos"]["x"] = i
                    sentence["pos"]["y"] = k
                    # 将处理结果加入self.sentences
                    self.sentences.append(sentence)
                    text = ""
                    k = k + 1

        for sentence in self.sentences:
            sentence["text"] = sentence["text"].strip()
            if sentence["text"] == "":
                self.sentences.remove(sentence)

        # 对文章位置进行标注，通过mark列表，标注出是否是第一段、尾段、第一句、最后一句
        lastpos = dict()
        lastpos["x"] = 0
        lastpos["y"] = 0
        lastpos["mark"] = list()
        for sentence in self.sentences:
            pos = sentence["pos"]
            pos["mark"] = list()
            if pos["x"] == 0:
                pos["mark"].append("FIRSTSECTION")
            if pos["y"] == 0:
                pos["mark"].append("FIRSTSENTENCE")
                lastpos["mark"].append("LASTSENTENCE")
            if pos["x"] == self.sentences[len(self.sentences)-1]["pos"]["x"]:
                pos["mark"].append("LASTSECTION")
            lastpos = pos
        lastpos["mark"].append("LASTSENTENCE")

    def __calc_keywords(self):
        # 计算tf-idfs，取出排名靠前的20个词
        words_best = list()
        words_best = words_best + jieba.analyse.extract_tags(self.text, topK=20)
        # words_best_textrank = jieba.analyse.textrank(self.text)  # by Lowenve
        # words_best = words_best + words_best_textrank  # by Lowenve
        # print(words_best_textrank)
        # for kw_textrank in words_best_textrank:
        #     if kw_textrank not in words_best:
        #         words_best.append(kw_textrank)
        # print(words_best)

        # 提取第一段的关键词
        parts = self.text.lstrip().split("\n")
        firstpart = ""
        """
        # 如果文章有分段，那么，找出“正文的第一段”赋值给：firstpart
        # 如何找出“文中的第一段”呢？
        #   - 先用回车符进行分段；
        #   - 再按段遍历文章，把第一个文字数量大于27的段落当做第一段；
        """
        for part in parts:
            if len(parts) >= 1 and len(part) > 60:
                firstpart = firstpart + part

        words_best = words_best + jieba.analyse.extract_tags(firstpart, topK=5)
        # 提取title中的关键词
        words_best = words_best + jieba.analyse.extract_tags(self.title, topK=3)
        # 将结果合并成一个句子，并进行分词
        text = ""
        for w in words_best:
            text = text + " " + w
        # 计算词性，提取名词和动词
        words = jieba.posseg.cut(text)
        keywords = list()
        for w in words:
            flag = w.flag
            word = w.word
            if flag.find('n') >= 0 or flag.find('v') >= 0:
                if len(word) > 1:
                    keywords.append(word)
        # 保留前20个关键词
        keywords = jieba.analyse.extract_tags(" ".join(keywords), topK=20)
        keywords = list(set(keywords))
        self.keywords = keywords

    def __calc_sentence_weight_by_keywords(self):
        # 计算句子的关键词权重
        for sentence in self.sentences:
            sentence["weightKeywords"] = 0
        for keyword in self.keywords:
            for sentence in self.sentences:
                if sentence["text"].find(keyword) >= 0:
                    sentence["weightKeywords"] = sentence["weightKeywords"] + 1

    def __calc_sentence_weight_by_pos(self):
        # 计算句子的位置权重
        for sentence in self.sentences:
            mark = sentence["pos"]["mark"]
            weight_pos = 0
            if "FIRSTSECTION" in mark:
                weight_pos = weight_pos + 2
            if "FIRSTSENTENCE" in mark:
                weight_pos = weight_pos + 2
            if "LASTSENTENCE" in mark:
                weight_pos = weight_pos + 1
            if "LASTSECTION" in mark:
                weight_pos = weight_pos + 1
            sentence["weight_pos"] = weight_pos

    def __calc_sentence_weight_by_cue_words(self):
        # 计算句子的线索词权重
        for sentence in self.sentences:
            sentence["weightCueWords"] = 0
        for i in CLUE_KEYWORDS_LIST:
            for sentence in self.sentences:
                if sentence["text"].find(i) >= 0:
                    sentence["weightCueWords"] = 1

    def __calc_sentence_weight(self):
        self.__calc_sentence_weight_by_pos()
        self.__calc_sentence_weight_by_cue_words()
        self.__calc_sentence_weight_by_keywords()
        for sentence in self.sentences:
            sentence["weight"] = sentence["weight_pos"] + 2 * sentence["weightCueWords"] + sentence["weightKeywords"]

    def get_summary(self, ratio=0.2):
        # 清空变量
        self.keywords = list()
        self.sentences = list()
        self.summary = list()

        # 调用方法，分别计算关键词、分句，计算权重
        self.__calc_keywords()
        self.__split_sentence()
        self.__calc_sentence_weight()

        # 对句子的权重值进行排序
        self.sentences = sorted(self.sentences, key=lambda k: k['weight'], reverse=True)

        if len(self.text) < 300:
            # 当信件内容少于300字符时，直接返回原文；
            for i in range(len(self.sentences)):
                if i < len(self.sentences):
                    sentence = self.sentences[i]
                    self.summary.append(sentence['text'])

        else:
            # 当信件内容很多，则根据排序结果，取排名占前X%的句子作为摘要
            for i in range(len(self.sentences)):
                if i < ratio * len(self.sentences):
                    sentence = self.sentences[i]
                    self.summary.append(sentence['text'])
        # logger.debug(msg="提取文本概要信息...")
        return self.summary

    def get_keywords(self):
        self.keywords = list()
        self.__calc_keywords()
        return self.keywords



class Likelihood:
    def word2vec(self, word1, word2):
        if self.punctuation is False:
            seg_list_1 = [w for w in list(jieba.cut(word1, cut_all=False)) if w not in PUN_LIST]
            seg_list_2 = [w for w in list(jieba.cut(word2, cut_all=False)) if w not in PUN_LIST]
        else:
            seg_list_1 = list(jieba.cut(word1, cut_all=False))
            seg_list_2 = list(jieba.cut(word2, cut_all=False))

        total_seg_list = list(set(seg_list_1 + seg_list_2))
        seg_vec_1 = []
        seg_vec_2 = []
        for word_tol in total_seg_list:
            freq = 0
            for word in seg_list_1:
                if word_tol == word:
                    freq += 1
            seg_vec_1.append(freq)
            freq = 0
            for word in seg_list_2:
                if word_tol == word:
                    freq += 1
            seg_vec_2.append(freq)
        self.seg_vec_1, self.seg_vec_2 = seg_vec_1, seg_vec_2

    def cos_dist(self):
        if len(self.seg_vec_1) != len(self.seg_vec_2):
            return None
        part_up = 0.0
        a_sq = 0.0
        b_sq = 0.0
        for a1, b1 in zip(self.seg_vec_1, self.seg_vec_2):
            part_up += a1 * b1
            a_sq += a1 ** 2
            b_sq += b1 ** 2
        part_down = math.sqrt(a_sq * b_sq)
        if part_down == 0.0:
            # return None
            return 0.0
        else:
            return part_up / part_down

    def likelihood(self, word1, word2, punctuation=False):
        self.word1 = word1
        self.word2 = word2
        self.punctuation = punctuation
        self.word2vec(self.word1, self.word2)
        like_per = self.cos_dist()
        return like_per

    def __split_sentence(self, letter_text):
        # 通过换行符对文档进行分段
        sections = letter_text.split("\n")
        for section in sections:
            if section == "":
                sections.remove(section)

        # 通过分割符对每个段落进行分句
        sentence_id = 0
        sentences_dict = dict()
        for i in range(len(sections)):
            section = sections[i]
            text = ""
            # k = 0
            for j in range(len(section)):
                char = section[j]
                text = text + char
                if char in ["!",  "。", "？", "；", "?", ";", "，", ",", ":", "："] or j == len(section)-1:
                    text = text.strip()
                    if len(text) > 0:
                        sentences_dict[sentence_id] = text
                    text = ""
                    sentence_id = sentence_id + 1

        return sentences_dict

    def compare(self, text1, text2):
        sentences1_dict = self.__split_sentence(letter_text=text1)
        sentences2_dict = self.__split_sentence(letter_text=text2)
        xiangtongjvzi = 0
        sum_jvzi = 0
        xiangtongjvzi_list = list()
        for key1, stc1 in sentences1_dict.items():
            sum_jvzi += 1
            highest_similarity = 0
            highest_similarity_sentence_id = 0
            # for key2, stc2 in sentences2_dict.items():
            #     similarity = self.likelihood(word1=stc1, word2=stc2)
            #     if similarity > highest_similarity:
            #         highest_similarity = similarity
            #         highest_similarity_sentence_id = key2
            # print(stc1)
            if len(stc1) > 7 and stc1 in text2:
                xiangtongjvzi += 1
            # sentences2_dict[highest_similarity_sentence_id] = "<span style='color:red'>" + sentences2_dict[
            #     highest_similarity_sentence_id] + '[' + str(key1) + "]</span>"
            # sentences1_dict[key1] = "<span style='color:red'>" + sentences1_dict[key1] + '[' + str(
            #     key1) + "]</span>"
                xiangtongjvzi_list.append(stc1)
            
            # elif highest_similarity > 0.6:
            #     sentences2_dict[highest_similarity_sentence_id] = "<span style='color:green'>" + sentences2_dict[
            #         highest_similarity_sentence_id] + '[' + str(key1) + "]</span>"
            #     sentences1_dict[key1] = "<span style='color:green'>" + sentences1_dict[key1] + '[' + str(
            #         key1) + "]</span>"

       
        bilv = xiangtongjvzi/sum_jvzi
        return bilv


def cal_similar_letter(keyword_letter: list):
    # 查找相似文本；
    all_letter_list = Projects.objects.all()
    repeat_letter_id_dict = {}
    for letter in all_letter_list:
        letter_keyword_repeat_num = 0
        for keyword in keyword_letter:
            if letter.key_word and keyword in letter.key_word:
                letter_keyword_repeat_num = letter_keyword_repeat_num + 1

        # 当相同的关键词数量至少为3个的时候，默认为是相似信件；
        if letter_keyword_repeat_num > 3:
            repeat_letter_id_dict[letter.id] = letter_keyword_repeat_num
    repeat_letter_id_dict = sorted(repeat_letter_id_dict.items(), key=lambda item: item[1], reverse=True)
    repeat_letter_id_list = []
    for letter_id in repeat_letter_id_dict[1:2]:
        repeat_letter_id_list.append(letter_id[0])

    repeat_letter_list = Projects.objects.filter(id__in=repeat_letter_id_list)
    return repeat_letter_list




def fuck_docx(sec):
    import os
    import re
    import docx
    from docx.shared import RGBColor#设置字体
    from docx import Document
    from docx.shared import Pt#设置字体
    from docx.oxml.ns import qn#设置中文字体
    f = open("test.txt","r")#将文档和python放在一个目录下无需复杂的路径
    i=0
    document = Document()#新建word
    p = document.add_paragraph('')#新建段落，这句话放在循环外面可以减少空行
    while i<3190:
        content = f.readline()
        #print(content)
        
        if content.find('答案')!=-1: #如果该行有关键字“答案”就以关键字为分界进行分割
            pt = r'(答案)'
            res = re.split(pt, content)#分割
            #print(res[0],end=' ')
            #print(res[1],end=' ')
            #print(res[2])
            
            #add_run在同一段添加内容
            run = p.add_run(res[0])#输入关键字之前的字符
            run.font.name=u'宋体' #设置插入的字体
            run.font.size = Pt(15)
            r = run._element
            r.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')


            run = p.add_run(res[1])#输入关键字
            run.font.name=u'宋体'  #设置插入的字体
            run.font.size = Pt(15)
            r = run._element
            r.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')

            run = p.add_run(res[2])#输入关键字之后的字符
            run.font.name=u'宋体'  #设置插入的字体
            run.font.size = Pt(15)
            r = run._element
            r.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')

            #设置关键字之后也就是答案的字体颜色，这里设置为白色
            run.font.color.rgb = RGBColor(250,0,0)
        else:


            run = p.add_run(content)#如果该行没有关键字“答案”则直接输入word
            #print(content)
            run.font.name=u'宋体'  #设置插入的字体
            run.font.size = Pt(15)
            r = run._element
            r.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')

        i+=1
    document.save('路线1.docx')#关闭保存word
    f.close() #关闭TXT